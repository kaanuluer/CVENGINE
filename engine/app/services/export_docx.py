from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.schemas import Resume, ResumeLanguage, TemplateName
from app.services.ats import headings_for, _fmt_date
from app.services.language import language_upper
from app.services.sanitize import sanitize_resume

NAVY = RGBColor(0x11, 0x18, 0x27)
GRAY = RGBColor(0x4B, 0x55, 0x63)
BLUE = RGBColor(0x3B, 0x82, 0xF6)


def _spacing(template: TemplateName) -> tuple[float, float, int]:
    if template == "executive":
        return 1.0, 0.9, 12
    if template == "compact":
        return 0.6, 0.6, 10
    return 0.75, 0.75, 11


def build_docx(resume: Resume, template: TemplateName = "classic", language: ResumeLanguage = "en") -> bytes:
    resume = sanitize_resume(resume)
    margin, _, body = _spacing(template)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(margin)
    section.bottom_margin = Inches(margin)
    section.left_margin = Inches(margin)
    section.right_margin = Inches(margin)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(body)
    style.font.color.rgb = NAVY
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    b = resume.basics
    h = headings_for(language)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = name.add_run(b.name)
    run.bold = True
    run.font.size = Pt(22 if template == "executive" else 18 if template == "classic" else 16)
    run.font.color.rgb = NAVY
    name.paragraph_format.space_after = Pt(2)

    if b.label:
        lab = doc.add_paragraph()
        r = lab.add_run(b.label)
        r.font.size = Pt(body + 1)
        r.font.color.rgb = BLUE if template == "executive" else GRAY
        lab.paragraph_format.space_after = Pt(2)

    contact = "  ·  ".join(p for p in [b.email, b.phone, b.url, b.location.city] if p)
    if contact:
        p = doc.add_paragraph()
        r = p.add_run(contact)
        r.font.size = Pt(body - 1)
        r.font.color.rgb = GRAY
        p.paragraph_format.space_after = Pt(8)

    def add_heading(title: str) -> None:
        p = doc.add_paragraph()
        r = p.add_run(language_upper(title, language))
        r.bold = True
        r.font.size = Pt(body)
        r.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(10 if template != "compact" else 6)
        p.paragraph_format.space_after = Pt(4)
        p_pr = p._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "3B82F6" if template == "executive" else "D1D5DB")
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

    def add_body(text: str, bold: bool = False) -> None:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(body)
        r.font.color.rgb = NAVY
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.08 if template == "executive" else 1.02

    def add_bullet(text: str) -> None:
        p = doc.add_paragraph(text, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(body)
            run.font.color.rgb = NAVY
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.2)

    if b.summary:
        add_heading(h["summary"])
        add_body(b.summary)

    if resume.work:
        add_heading(h["experience"])
        for work in resume.work:
            dates = " – ".join(
                p
                for p in [
                    _fmt_date(work.startDate, language),
                    _fmt_date(work.endDate, language)
                    or ("Present" if language == "en" and work.startDate else ("Günümüz" if work.startDate else "")),
                ]
                if p
            )
            left = "  ·  ".join(p for p in [work.position, work.name] if p)
            line = f"{left}    {dates}".strip()
            add_body(line, bold=True)
            for item in work.highlights:
                add_bullet(item)

    if resume.education:
        add_heading(h["education"])
        for edu in resume.education:
            left = "  ·  ".join(p for p in [edu.studyType, edu.area, edu.institution] if p)
            dates = " – ".join(p for p in [_fmt_date(edu.startDate, language), _fmt_date(edu.endDate, language)] if p)
            add_body(f"{left}    {dates}".strip(), bold=True)

    if resume.skills:
        add_heading(h["skills"])
        for skill in resume.skills:
            label = f"{skill.name}: " if skill.name else ""
            add_body(label + ", ".join(skill.keywords))

    if resume.projects:
        add_heading(h["projects"])
        for project in resume.projects:
            add_body(project.name, bold=True)
            if project.description:
                add_body(project.description)
            for item in project.highlights:
                add_bullet(item)

    if resume.certificates:
        add_heading(h["certificates"])
        for cert in resume.certificates:
            add_body("  ·  ".join(p for p in [cert.name, cert.issuer] if p))

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def write_docx(path: Path, resume: Resume, template: TemplateName = "classic", language: ResumeLanguage = "en") -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(build_docx(resume, template, language))
    return path
