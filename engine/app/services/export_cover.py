from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas

from app.services.export_pdf import FONTS

NAVY = HexColor("#111827")
DOCX_NAVY = RGBColor(0x11, 0x18, 0x27)


def build_cover_pdf(text: str) -> bytes:
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    x0, y = 22 * mm, height - 22 * mm
    max_width = width - 44 * mm
    c.setFillColor(NAVY)
    for paragraph in _paragraphs(text):
        y = _draw_paragraph(c, paragraph, x0, y, max_width)
        y -= 12
        if y < 28 * mm:
            c.showPage()
            y = height - 22 * mm
            c.setFillColor(NAVY)
    c.save()
    return buffer.getvalue()


def build_cover_docx(text: str) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = DOCX_NAVY
    for paragraph in _paragraphs(text):
        p = doc.add_paragraph()
        lines = paragraph.split("\n")
        for i, line in enumerate(lines):
            run = p.add_run(line)
            run.font.size = Pt(11)
            run.font.color.rgb = DOCX_NAVY
            if i < len(lines) - 1:
                run.add_break()
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.15
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def write_cover_pdf(path: Path, text: str) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(build_cover_pdf(text))
    return path


def write_cover_docx(path: Path, text: str) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(build_cover_docx(text))
    return path


def _paragraphs(text: str) -> list[str]:
    return [block.strip() for block in text.replace("\r\n", "\n").split("\n\n") if block.strip()]


def _draw_paragraph(c, text: str, x0: float, y: float, max_width: float) -> float:
    font, size, leading = FONTS["reg"], 11, 15
    c.setFont(font, size)
    c.setFillColor(NAVY)
    for raw_line in text.split("\n"):
        if not raw_line.strip():
            y -= leading * 0.4
            continue
        line = ""
        for word in raw_line.split():
            trial = (line + " " + word).strip()
            if c.stringWidth(trial, font, size) <= max_width:
                line = trial
            else:
                c.drawString(x0, y, line)
                y -= leading
                line = word
        if line:
            c.drawString(x0, y, line)
            y -= leading
    return y
